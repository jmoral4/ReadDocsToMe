# Code
import hashlib
import json
import pygame
import argparse
from pathlib import Path
from openai import OpenAI
from docx import Document
from halo import Halo

# ANSI escape codes for colors
RED = "\033[31m"
GREEN = "\033[32m"
YELLOW = "\033[33m"
BLUE = "\033[34m"
MAGENTA = "\033[35m"
CYAN = "\033[36m"
WHITE = "\033[37m"
RESET = "\033[0m"

spinner = Halo(spinner='dots')

def print_colored(text, color):
    print(f"{color}{text}{RESET}")


def play_mp3(filepath):
    try:
        pygame.mixer.init()
        pygame.mixer.music.load(filepath)
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)
    except pygame.error as e:
        print_colored(f"Error playing MP3 {filepath}: {e}", RED)
    finally:
        pygame.mixer.quit()

@Halo(text='Generating Audio', spinner='dots')
def generate_audio(content, speech_file_path, voice="nova"):
    client = OpenAI(api_key=API_KEY)
    try:
        audio_resp = client.audio.speech.create(
            model="gpt-4o-mini-tts",
            voice=voice,
            input=str(content) # Ensure input is a string
        )
        # The response object has a .write_to_file() method
        audio_resp.write_to_file(speech_file_path)
    except Exception as e:
        print_colored(f"Error during OpenAI API call or file writing: {e}", RED)

        if speech_file_path.exists():
            try:
                speech_file_path.unlink() # Remove potentially corrupted/empty file
            except OSError as oe:
                print_colored(f"Could not remove partial file {speech_file_path}: {oe}", RED)


def read_word_document(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        # Get text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Get text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(' | '.join(row_text))

        return '\n'.join(full_text)
    except Exception as e:
        print_colored(f"Error reading Word document {file_path}: {e}", RED)
        return None

def read_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print_colored(f"Error reading text file {file_path}: {e}", RED)
        return None

def chunk_text(text, max_chars=4050):  # API limit is 4096 for TTS
    words = text.split()
    chunks = []
    current_chunk_list = []
    current_length = 0

    for word in words:
        if not current_chunk_list:
            current_chunk_list.append(word)
            current_length = len(word)
        elif current_length + len(word) + 1 > max_chars:
            chunks.append(" ".join(current_chunk_list))
            current_chunk_list = [word]
            current_length = len(word)
        else:
            current_chunk_list.append(word)
            current_length += len(word) + 1

    if current_chunk_list:
        chunks.append(" ".join(current_chunk_list))

    return chunks

def generate_file_hash(file_path):
    try:
        with open(file_path, 'rb') as file:
            return hashlib.md5(file.read()).hexdigest()
    except FileNotFoundError:
        print_colored(f"File not found for hashing: {file_path}", YELLOW)
        return None
    except Exception as e:
        print_colored(f"Error generating hash for {file_path}: {e}", RED)
        return None


def process_document(doc_path, output_dir, force_regenerate=False, voice="nova", silent=False, fixed_filename=None):
    if not silent:
        print_colored("Getting content...", BLUE)


    doc_hash = generate_file_hash(doc_path)
    if doc_hash is None and not force_regenerate: # If hashing failed and not forcing, maybe skip
        print_colored(f"Could not generate hash for {doc_path}. To process, use --force.", RED)
        return None

    doc_stem = fixed_filename if fixed_filename else Path(doc_path).stem
    hash_file = Path(output_dir) / f"{doc_stem}_hash.txt"
    # Ensure output directory exists
    Path(output_dir).mkdir(parents=True, exist_ok=True)


    if hash_file.exists() and not force_regenerate:
        try:
            with open(hash_file, 'r') as f:
                saved_hash = f.read().strip()
            if saved_hash == doc_hash:
                print_colored("Document unchanged, using existing audio files.", GREEN)
                return Path(output_dir) / doc_stem
        except Exception as e:
            print_colored(f"Error reading hash file {hash_file}: {e}. Regenerating.", YELLOW)


    if str(doc_path).endswith('.docx'):
        content = read_word_document(doc_path)
    else:
        content = read_text_file(doc_path)

    if content is None or not content.strip():
        print_colored(f"Failed to read document content or content is empty for {doc_path}.", RED)
        return None

    chunks = chunk_text(content)
    if not chunks:
        print_colored(f"No text chunks generated from {doc_path}. Document might be empty or too short.", YELLOW)
        return None
    print_colored(f"Document split into {len(chunks)} chunks", BLUE)

    if not silent:
        play_mp3('genaudio.mp3')
        print_colored("Generating audio...", BLUE)

    success_count = 0
    generated_files_base = Path(output_dir) / doc_stem

    for i, chunk in enumerate(chunks):
        chunk_file_path = Path(output_dir) / f"{doc_stem}_{i+1}.mp3"
        # Check if chunk file already exists and if we should skip (relevant if a previous run failed mid-way)
        if chunk_file_path.exists() and not force_regenerate:
            # A more robust check would be to hash the chunk content and compare,
            # but for simplicity, we assume if the file exists and not forcing, it's good.
            # This part is tricky if the content has changed but the number of chunks is the same.
            # The top-level document hash check should ideally cover this.
            print_colored(f"Audio chunk {i+1}/{len(chunks)} ({chunk_file_path.name}) already exists. Skipping generation.", CYAN)
            success_count += 1
            continue

        try:
            print_colored(f"Generating audio chunk {i+1}/{len(chunks)} for {chunk_file_path.name}", YELLOW)
            generate_audio(chunk, chunk_file_path, voice) # generate_audio now takes Path object
            if chunk_file_path.exists() and chunk_file_path.stat().st_size > 0: # Check if file was actually created
                print_colored(f"Generated: {chunk_file_path}", GREEN)
                success_count += 1
            else:
                print_colored(f"Audio file {chunk_file_path} was not created or is empty for chunk {i+1}.", RED)

        except Exception as e: # Catch exceptions from generate_audio if it re-raises
            print_colored(f"Failed to generate audio for chunk {i+1}: {e}", RED)

    if success_count == len(chunks):
        if doc_hash: # Only write hash if we successfully got one
            with open(hash_file, 'w') as f:
                f.write(doc_hash)
            print_colored(f"All {len(chunks)} chunks generated successfully. Hash saved.", GREEN)
    else:
        print_colored(f"Only {success_count}/{len(chunks)} chunks generated successfully. Hash not updated.", YELLOW)
        if hash_file.exists() and doc_hash: # If some failed, the old hash might be invalid
             print_colored(f"Consider removing {hash_file} if you want to ensure regeneration on next run without --force", YELLOW)


    return generated_files_base


def play_audio_sequence(base_path):
    audio_files = sorted(list(Path(base_path.parent).glob(f"{base_path.name}_*.mp3")))

    if not audio_files:
        print_colored(f"No audio files found with base name {base_path.name}", RED)
        return

    spinner.text = 'Now Playing'
    spinner.start()

    for i, audio_file in enumerate(audio_files):
        print_colored(f"Playing chunk {i+1}/{len(audio_files)}: {audio_file}", CYAN)
        play_mp3(str(audio_file))

    spinner.stop()

def process_document_folder(folder_path, output_dir, force_regenerate=False, voice="nova", silent=False, fixed_filename=None):
    folder = Path(folder_path)
    if not folder.is_dir():
        print_colored(f"{folder_path} is not a directory", RED)
        return

    doc_files = list(folder.glob("*.docx")) + list(folder.glob("*.txt"))
    if not doc_files:
        print_colored("No Word documents (.docx) or text files (.txt) found in the folder", RED)
        return

    for idx, doc_file in enumerate(doc_files):
        print_colored(f"Processing: {doc_file}", BLUE)
        current_fixed_filename = fixed_filename
        if fixed_filename and len(doc_files) > 1:
            current_fixed_filename = f"{fixed_filename}_{idx + 1}"

        base_path = process_document(doc_file, output_dir, force_regenerate, voice, silent, current_fixed_filename)
        if base_path and not args.download_only:
            # Check if any audio files were actually produced for this base_path
            # This is important if process_document returned a base_path but failed to generate all/any chunks
            actual_audio_files = sorted(list(Path(base_path.parent).glob(f"{base_path.name}_*.mp3")))
            if actual_audio_files:
                play_audio_sequence(base_path)
            else:
                print_colored(f"No audio files to play for {base_path.name}", YELLOW)


API_KEY = None

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Document to Podcast Converter")
    parser.add_argument("--document", help="Path to Word document or text file to convert", default=None)
    parser.add_argument("--folder", help="Path to folder containing Word documents or text files to convert", default=None)
    parser.add_argument("--output-dir", help="Directory to save the generated audio files", default=None)
    parser.add_argument("--force", help="Force regeneration of audio even if document hasn't changed", action='store_true', default=False)
    parser.add_argument("--download-only", help="Only generate audio files, no playback", action='store_true', default=False)
    parser.add_argument("--voice", help="Voice to use for TTS (alloy, echo, fable, onyx, nova, shimmer)", default=None)
    parser.add_argument("--silent", help="Don't vocalize actions (placeholder sounds) being performed", action='store_true', default=False)
    parser.add_argument("--fixed-filename", help="Use a fixed filename prefix for the audio output", default=None)

    args = parser.parse_args()

    try:
        with open('config.json') as config_file:
            config = json.load(config_file)
            API_KEY = config['OPENAI_KEY']
            OUTPUT_DIR = args.output_dir if args.output_dir else config.get('OUTPUT_DIR', 'audio_output') # Provide default if not in config
            AUDIO_VOICE = args.voice if args.voice else config.get('AUDIO_VOICE', 'nova')
    except FileNotFoundError:
        print_colored("Error: config.json not found. Please create it with your OPENAI_KEY.", RED)
        exit(1)
    except KeyError as e:
        print_colored(f"Error: Missing key {e} in config.json.", RED)
        exit(1)

    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    # 2.0 since we're deriving from 'ReadItToMe' codebase - though nearly entirely changed. >90%
    print_colored("Document to Podcast Converter 2.0 (using gpt-4o-mini-tts)", MAGENTA)

    # Test doc -- 50 page DeepResearch doc
    # args.document = r"C:\Docs\Launching a Twitch.docx"

    if args.document:
        doc_path = Path(args.document)
        if not doc_path.exists():
            print_colored(f"Document not found: {doc_path}", RED)
        else:
            base_path = process_document(doc_path, OUTPUT_DIR, args.force, AUDIO_VOICE, args.silent, args.fixed_filename)
            if base_path and not args.download_only:
                # Check if any audio files were actually produced
                actual_audio_files = sorted(list(Path(base_path.parent).glob(f"{base_path.name}_*.mp3")))
                if actual_audio_files:
                    play_audio_sequence(base_path)
                else:
                    print_colored(f"No audio files to play for {base_path.name}", YELLOW)
    elif args.folder:
        process_document_folder(args.folder, OUTPUT_DIR, args.force, AUDIO_VOICE, args.silent, args.fixed_filename)
    else:
        print_colored("Please provide either a document path (--document) or a folder path (--folder)", RED)
        parser.print_help()

    print_colored("All Done!", GREEN)